"""
R22: DOCX 내보내기 목차(TOC) 자동 생성 테스트
"""
from __future__ import annotations

import unittest
import unittest.mock
from unittest.mock import patch

from routes.export_routes import _extract_headings, _insert_toc


class TestExtractHeadings(unittest.TestCase):
    """마크다운에서 헤딩 추출 테스트"""

    def test_basic_headings(self):
        md = "# 제목1\n내용\n## 소제목1\n내용\n### 하위\n## 소제목2"
        headings = _extract_headings(md)
        self.assertEqual(len(headings), 4)
        self.assertEqual(headings[0], (1, '제목1'))
        self.assertEqual(headings[1], (2, '소제목1'))
        self.assertEqual(headings[2], (3, '하위'))
        self.assertEqual(headings[3], (2, '소제목2'))

    def test_no_headings(self):
        md = "일반 텍스트만 있는 콘텐츠입니다.\n또 다른 줄."
        headings = _extract_headings(md)
        self.assertEqual(headings, [])

    def test_strips_markdown_formatting(self):
        md = "## **굵은 제목**\n## *기울임*\n## `코드`"
        headings = _extract_headings(md)
        self.assertEqual(headings[0], (2, '굵은 제목'))
        self.assertEqual(headings[1], (2, '기울임'))
        self.assertEqual(headings[2], (2, '코드'))

    def test_ignores_hash_without_space(self):
        """#뒤에 바로 텍스트가 오는 경우도 헤딩으로 인식 (마크다운 호환)"""
        md = "#제목\n##제목2"
        headings = _extract_headings(md)
        # '#제목' → level=1, text='제목'
        self.assertEqual(len(headings), 2)

    def test_level_4_heading(self):
        md = "#### 4단계 헤딩"
        headings = _extract_headings(md)
        self.assertEqual(headings[0], (4, '4단계 헤딩'))

    def test_level_5_ignored(self):
        """5단계 이상 헤딩은 무시"""
        md = "##### 5단계"
        headings = _extract_headings(md)
        self.assertEqual(headings, [])


class TestInsertToc(unittest.TestCase):
    """DOCX 목차 삽입 테스트"""

    def test_toc_creates_paragraphs(self):
        from docx import Document
        doc = Document()
        md = "## 소개\n## 본문\n### 세부 내용\n## 결론"
        _insert_toc(doc, md)
        # 목차 제목(heading) + 4개 항목 + 구분선 = 6개 요소
        # Document 기본 빈 단락 포함
        texts = [p.text for p in doc.paragraphs]
        self.assertIn('목차', texts)
        # 항목이 포함되어 있는지
        combined = '\n'.join(texts)
        self.assertIn('소개', combined)
        self.assertIn('본문', combined)
        self.assertIn('세부 내용', combined)
        self.assertIn('결론', combined)

    def test_toc_empty_content(self):
        """헤딩이 없으면 목차를 삽입하지 않음"""
        from docx import Document
        doc = Document()
        initial_count = len(doc.paragraphs)
        _insert_toc(doc, "일반 텍스트만")
        self.assertEqual(len(doc.paragraphs), initial_count)


class TestExportDocxWithToc(unittest.TestCase):
    """/api/export/docx 라우트의 include_toc 파라미터 테스트"""

    def setUp(self):
        from app import create_app
        self.app = create_app()
        self.client = self.app.test_client()

    @patch('services.data.supabase_service.is_supabase_enabled', return_value=False)
    def test_export_with_toc(self, _):
        """include_toc=true 시 DOCX에 목차 포함"""
        resp = self.client.post('/api/export/docx',
            json={
                'title': '테스트 문서',
                'content': '## 소개\n내용1\n## 결론\n내용2',
                'include_toc': True,
            },
            headers={'Origin': 'http://localhost:3000'}
        )
        self.assertEqual(resp.status_code, 200)
        self.assertIn('application/vnd.openxmlformats', resp.content_type)
        # DOCX 바이너리를 파싱하여 목차 확인
        import io
        from docx import Document
        doc = Document(io.BytesIO(resp.data))
        texts = [p.text for p in doc.paragraphs]
        combined = '\n'.join(texts)
        self.assertIn('목차', combined)
        self.assertIn('소개', combined)

    @patch('services.data.supabase_service.is_supabase_enabled', return_value=False)
    def test_export_without_toc(self, _):
        """include_toc 없으면 목차 미포함"""
        resp = self.client.post('/api/export/docx',
            json={
                'title': '테스트 문서',
                'content': '## 소개\n내용1',
            },
            headers={'Origin': 'http://localhost:3000'}
        )
        self.assertEqual(resp.status_code, 200)
        import io
        from docx import Document
        doc = Document(io.BytesIO(resp.data))
        texts = [p.text for p in doc.paragraphs]
        self.assertNotIn('목차', texts)


if __name__ == '__main__':
    unittest.main()
