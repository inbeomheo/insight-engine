"""
DOCX 내보내기 서비스
마크다운 콘텐츠를 Word 문서로 변환
"""
import io
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# 한글 폰트 설정 (Windows: Malgun Gothic, Linux: Noto Sans KR)
def _get_font_name():
    try:
        import platform
        if platform.system() == 'Windows':
            return 'Malgun Gothic'
    except Exception:
        pass
    return 'Noto Sans KR'


DEFAULT_FONT = _get_font_name()


def markdown_to_docx(title: str, content: str) -> io.BytesIO:
    """마크다운 텍스트를 DOCX 바이너리로 변환합니다.

    Args:
        title: 문서 제목
        content: 마크다운 본문

    Returns:
        BytesIO: DOCX 파일 바이너리 스트림
    """
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = DEFAULT_FONT
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # 제목 추가
    if title:
        heading = doc.add_heading(title, level=0)
        for run in heading.runs:
            run.font.name = DEFAULT_FONT
            run.font.size = Pt(20)

    # 마크다운 파싱 및 변환
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # 빈 줄
        if not line.strip():
            i += 1
            continue

        # 헤딩
        heading_match = re.match(r'^(#{1,6})\s+(.+)', line)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            text = heading_match.group(2).strip()
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.name = DEFAULT_FONT
            i += 1
            continue

        # 인용문 (>)
        if line.strip().startswith('>'):
            quote_text = line.strip().lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(quote_text)
            run.font.name = DEFAULT_FONT
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            i += 1
            continue

        # 불릿 리스트 (-, *, +)
        list_match = re.match(r'^[\s]*[-*+]\s+(.+)', line)
        if list_match:
            text = list_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)
            i += 1
            continue

        # 순서 리스트 (1. 2. 3.)
        ol_match = re.match(r'^[\s]*\d+[.)]\s+(.+)', line)
        if ol_match:
            text = ol_match.group(1)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, text)
            i += 1
            continue

        # 구분선 (---, ***, ___)
        if re.match(r'^[-*_]{3,}\s*$', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run('─' * 50)
            run.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
            run.font.size = Pt(8)
            i += 1
            continue

        # 일반 단락
        p = doc.add_paragraph()
        _add_formatted_text(p, line)
        i += 1

    # BytesIO로 저장
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _add_formatted_text(paragraph, text):
    """인라인 마크다운 서식을 적용하여 텍스트를 단락에 추가합니다.

    지원: **굵게**, *기울임*, `코드`, [링크](url)
    """
    # 인라인 패턴 매칭
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)'    # **bold**
        r'|(\*(.+?)\*)'       # *italic*
        r'|(`(.+?)`)'         # `code`
        r'|(\[(.+?)\]\((.+?)\))'  # [text](url)
    )

    last_end = 0
    for match in pattern.finditer(text):
        # 매치 이전 텍스트
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end:match.start()])
            run.font.name = DEFAULT_FONT

        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.font.name = DEFAULT_FONT
            run.font.bold = True
        elif match.group(4):  # *italic*
            run = paragraph.add_run(match.group(4))
            run.font.name = DEFAULT_FONT
            run.font.italic = True
        elif match.group(6):  # `code`
            run = paragraph.add_run(match.group(6))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
        elif match.group(8):  # [text](url)
            link_text = match.group(8)
            run = paragraph.add_run(link_text)
            run.font.name = DEFAULT_FONT
            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            run.font.underline = True

        last_end = match.end()

    # 나머지 텍스트
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = DEFAULT_FONT


def export_markdown(title: str, content: str) -> io.BytesIO:
    """마크다운 파일을 BytesIO로 반환합니다."""
    text = f"# {title}\n\n{content}"
    buffer = io.BytesIO(text.encode('utf-8'))
    return buffer


def export_txt(title: str, content: str) -> io.BytesIO:
    """순수 텍스트 파일을 BytesIO로 반환합니다. 마크다운 서식을 제거합니다."""
    text = re.sub(r'#{1,6}\s+', '', content)
    text = re.sub(r'\*{1,3}(.+?)\*{1,3}', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    full = f"{title}\n{'=' * len(title)}\n\n{text}"
    buffer = io.BytesIO(full.encode('utf-8'))
    return buffer


def export_zip(title: str, content: str) -> io.BytesIO:
    """DOCX + MD + TXT + meta.json을 ZIP으로 묶어 반환합니다."""
    import zipfile
    import json
    from datetime import datetime, timezone

    zip_buffer = io.BytesIO()
    safe_name = re.sub(r'[^\w\s가-힣-]', '', title)[:30].strip() or 'content'

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        # Markdown
        md_buf = export_markdown(title, content)
        zf.writestr(f'{safe_name}.md', md_buf.getvalue())

        # TXT
        txt_buf = export_txt(title, content)
        zf.writestr(f'{safe_name}.txt', txt_buf.getvalue())

        # DOCX
        docx_buf = markdown_to_docx(title, content)
        zf.writestr(f'{safe_name}.docx', docx_buf.getvalue())

        # meta.json
        meta = {
            'title': title,
            'char_count': len(content),
            'exported_at': datetime.now(timezone.utc).isoformat(),
            'formats': ['md', 'txt', 'docx'],
        }
        zf.writestr('meta.json', json.dumps(meta, ensure_ascii=False, indent=2))

    zip_buffer.seek(0)
    return zip_buffer
