"""
블로그 콘텐츠 생성 API 라우트
"""
import concurrent.futures
import io
import json
import re as re_module
import time
from typing import Dict

from flask import Blueprint, request, jsonify, current_app, g, Response, stream_with_context, send_file

from config import get_model_max_tokens
from services import ai_service, content_service, fusion_service
from services.content_service import clear_cache
from services.supabase_service import (
    require_auth, is_supabase_enabled, save_history
)
from services.usage import require_usage, check_usage
from services.usage.usage_decorator import get_usage_for_response
import uuid

blog_bp = Blueprint('blog', __name__)
_CLIENT_TRACKER: Dict[str, float] = {}

DEFAULT_MODEL = 'zhipuai/GLM-4.5-Air'
DEFAULT_STYLE = 'blog_seo'
MAX_BATCH_URLS = 10
MAX_BATCH_WORKERS = 5
BATCH_CONTENT_TOKEN_LIMIT = 3000


def _extract_client_id(req) -> str:
    """요청에서 클라이언트 ID를 추출합니다."""
    data = req.get_json(silent=True)
    if isinstance(data, dict) and data.get('clientId'):
        return str(data['clientId'])

    form_id = req.form.get('clientId')
    if form_id:
        return str(form_id)

    raw = (req.get_data(cache=False, as_text=True) or '').strip()
    if raw:
        try:
            parsed = json.loads(raw)
            if isinstance(parsed, dict) and parsed.get('clientId'):
                return str(parsed['clientId'])
        except json.JSONDecodeError:
            pass  # 잘못된 JSON은 무시하고 빈 문자열 반환

    return ''


def _get_request_data(req):
    """JSON 또는 form 데이터에서 공통 파라미터를 추출하고 검증합니다.
    API 키는 서버 환경변수에서 관리되므로 요청에서 추출하지 않습니다.
    """
    data = req.get_json(silent=True)
    if isinstance(data, dict) and data:
        # modifiers 검증
        modifiers, _ = _validate_modifiers(data.get('modifiers'))
        # custom_prompt 검증
        custom_prompt, _ = _validate_custom_prompt(data.get('customPrompt'))

        # urls 검증 (리스트 형식 확인)
        urls = data.get('urls', [])
        if not isinstance(urls, list):
            urls = []
        urls = [u for u in urls if isinstance(u, str)][:MAX_BATCH_URLS]

        return {
            'url': data.get('url') if isinstance(data.get('url'), str) else None,
            'urls': urls,
            'content': data.get('content') if isinstance(data.get('content'), str) else None,
            'model': data.get('model', DEFAULT_MODEL) if isinstance(data.get('model'), str) else DEFAULT_MODEL,
            'style': data.get('style', DEFAULT_STYLE) if isinstance(data.get('style'), str) else DEFAULT_STYLE,
            'modifiers': modifiers,
            'custom_prompt': custom_prompt,
        }

    return {
        'url': req.form.get('url'),
        'urls': [],
        'content': req.form.get('content'),
        'model': req.form.get('model', DEFAULT_MODEL),
        'style': req.form.get('style', DEFAULT_STYLE),
        'modifiers': None,
        'custom_prompt': None,
    }


def _validate_modifiers(modifiers):
    """modifiers 파라미터의 유효성을 검증합니다.

    Args:
        modifiers: dict 또는 None

    Returns:
        tuple: (validated_modifiers, error_message)
    """
    if modifiers is None:
        return None, None

    if not isinstance(modifiers, dict):
        return None, 'modifiers는 객체 형식이어야 합니다.'

    # 허용된 키와 값 정의 (v3.0: 2개 모디파이어만 지원)
    allowed_keys = {'length', 'writing_style'}
    allowed_values = {
        'length': {'short', 'medium', 'long'},
        'writing_style': {'conversational', 'explanatory', 'casual', 'expert'}
    }

    validated = {}
    for key, value in modifiers.items():
        if key not in allowed_keys:
            continue  # 알 수 없는 키는 무시
        if not isinstance(value, str):
            continue
        # 값 검증
        if key in allowed_values and value not in allowed_values[key]:
            continue  # 잘못된 값은 무시
        validated[key] = value[:50]  # 길이 제한

    return validated, None


def _validate_custom_prompt(custom_prompt):
    """customPrompt 파라미터의 유효성을 검증합니다.

    Returns:
        tuple: (validated_prompt, error_message)
    """
    if custom_prompt is None:
        return None, None

    if not isinstance(custom_prompt, str):
        return None, 'customPrompt는 문자열이어야 합니다.'

    # 길이 제한 (2000자)
    return custom_prompt.strip()[:2000], None


def _get_style_prompt(style, custom_prompt=None):
    """스타일에 맞는 프롬프트를 반환합니다."""
    if custom_prompt and custom_prompt.strip():
        return custom_prompt.strip()[:2000]
    style_prompts = current_app.config.get('STYLE_PROMPTS', {})
    return style_prompts.get(style, '')


def _handle_error_response(error_msg, log_detail=None):
    """에러 메시지에 따른 적절한 HTTP 상태 코드를 반환합니다.

    Args:
        error_msg: 사용자에게 보여줄 에러 메시지
        log_detail: 로깅용 상세 에러 (선택사항)
    """
    # 로깅 (상세 정보는 서버에만 기록)
    if log_detail:
        current_app.logger.error(f'Error: {log_detail}')

    # 세분화된 에러 메시지 (ai_service에서 이미 처리된 경우)는 그대로 전달
    error_prefixes = ['[인증 실패]', '[사용량 초과]', '[모델 오류]', '[타임아웃]',
                      '[연결 실패]', '[서비스 불가]', '[서버 오류]', '[잔액 부족]',
                      '[컨텐츠 차단]', '[AI 오류]']

    is_formatted_error = any(error_msg.startswith(prefix) for prefix in error_prefixes)

    if is_formatted_error:
        # 인증 관련 에러는 401
        if error_msg.startswith('[인증 실패]'):
            return jsonify({'error': error_msg}), 401
        # 사용량 초과는 429
        if error_msg.startswith('[사용량 초과]'):
            return jsonify({'error': error_msg}), 429
        # 나머지는 503 (서비스 일시 불가)
        return jsonify({'error': error_msg}), 503

    # API 키 관련 에러 (레거시)
    if 'API 키' in error_msg or 'authentication' in error_msg.lower():
        return jsonify({'error': error_msg}), 401

    safe_msg = _sanitize_error_for_client(error_msg)
    return jsonify({'error': safe_msg}), 500


def _sanitize_error_for_client(error_msg: str) -> str:
    """에러 메시지에서 내부 정보를 제거하여 클라이언트에 안전한 메시지를 반환합니다."""
    ALLOWED_ERROR_PREFIXES = [
        '[인증 실패]', '[사용량 초과]', '[입력 오류]', '[자막 없음]',
        '[생성 실패]', '[타임아웃]',
        '자막을', '댓글을', 'YouTube', 'API', '영상', 'URL'
    ]

    is_safe = any(error_msg.startswith(prefix) for prefix in ALLOWED_ERROR_PREFIXES)
    has_internal_info = any(keyword in error_msg.lower() for keyword in [
        'traceback', 'exception', 'file "', 'line ', 'error:', 'failed:',
        '/home/', '/usr/', 'supabase', 'postgres', 'connection',
        'api_key', 'token', 'secret', 'password', 'authorization', '.env',
        'openai', 'litellm', 'httpx', 'requests.exceptions',
    ])

    if not is_safe or has_internal_info:
        current_app.logger.error(f'Internal error hidden from user: {error_msg}')
        return '[서버 오류] 처리 중 예상치 못한 오류가 발생했습니다. 다시 시도해주세요.'

    return error_msg


def _fetch_youtube_content(video_id):
    """YouTube 영상의 자막과 댓글을 분리하여 가져옵니다.
    Supadata API 키는 환경변수에서 자동으로 로드됩니다.

    Returns:
        tuple: (transcript_text, comments_list, error, raw_transcript, transcript_source)
        - comments_list: 댓글 문자열 리스트 (전체) 또는 빈 리스트
        - transcript_source: 'api' | 'watch' | 'supadata' | 'cache'
    """
    transcript_result = content_service.get_transcript(video_id)
    if isinstance(transcript_result, dict) and transcript_result.get('error'):
        return None, [], transcript_result['error'], None, None

    # 새 형식: {'text': '...', 'source': '...'}
    transcript_text = transcript_result.get('text', '')
    transcript_source = transcript_result.get('source', 'unknown')

    comments = content_service.get_top_comments(video_id) or []

    return transcript_text, comments, None, transcript_text, transcript_source


def _build_combined_content(transcript_text, comments):
    """자막과 댓글을 기존 형식으로 합성합니다 (배치 처리 호환용).

    Args:
        transcript_text: 자막 텍스트
        comments: 댓글 리스트

    Returns:
        str: "[영상 자막]\\n...\\n\\n[시청자 댓글]\\n..." 형식 문자열
    """
    comments_text = '\n'.join(comments[:20]) if comments else '(댓글 없음)'
    return f"[영상 자막]\n{transcript_text}\n\n[시청자 댓글]\n{comments_text}"


@blog_bp.route('/health')
def health():
    """헬스체크 엔드포인트 (Railway/Docker용)"""
    return jsonify({'status': 'healthy'}), 200


@blog_bp.route('/')
def home():
    """API 서버 상태를 반환합니다. 프론트엔드는 Next.js에서 제공."""
    return jsonify({'status': 'ok', 'message': 'Insight Engine API Server'})


@blog_bp.route('/api/heartbeat', methods=['POST'])
def api_heartbeat():
    """클라이언트 연결 상태를 추적합니다."""
    client_id = _extract_client_id(request)
    if not client_id:
        return jsonify({'ok': False, 'error': 'clientId required'}), 400
    _CLIENT_TRACKER[client_id] = time.time()
    return jsonify({'ok': True})


@blog_bp.route('/api/close', methods=['POST'])
def api_close():
    """클라이언트 연결 종료를 처리합니다."""
    client_id = _extract_client_id(request)
    if not client_id:
        return jsonify({'ok': False, 'error': 'clientId required'}), 400
    _CLIENT_TRACKER.pop(client_id, None)
    return jsonify({'ok': True})


@blog_bp.route('/api/providers', methods=['GET'])
def api_providers():
    """API 키가 설정된 AI 서비스 및 모델 목록을 반환합니다.
    환경변수에 API 키가 설정된 프로바이더만 반환됩니다.
    """
    from config import get_available_providers, SUPADATA_API_KEY

    providers = get_available_providers()
    styles = current_app.config.get('STYLE_OPTIONS', [])

    return jsonify({
        'providers': providers,
        'styles': [{'id': s[0], 'name': s[1]} for s in styles],
        'supadataConfigured': bool(SUPADATA_API_KEY),
        'hasAutoFallback': True
    })


@blog_bp.route('/api/cache', methods=['DELETE'])
def api_clear_cache():
    """캐시를 삭제합니다. video_id 파라미터가 있으면 해당 영상만, 없으면 전체 삭제."""
    data = request.get_json(silent=True) or {}
    video_id = data.get('videoId')

    # URL에서 video_id 추출 (URL이 전달된 경우)
    url = data.get('url')
    if url and not video_id:
        video_id = content_service.get_video_id(url)

    deleted = clear_cache(video_id)

    if video_id:
        return jsonify({
            'success': True,
            'message': f'영상 {video_id}의 캐시가 삭제되었습니다.',
            'deleted': deleted
        })
    return jsonify({
        'success': True,
        'message': '전체 캐시가 삭제되었습니다.',
        'deleted': deleted
    })


@blog_bp.route('/api/recommend-style', methods=['POST'])
def recommend_style():
    """YouTube 제목을 분석하여 최적의 스타일과 모디파이어를 AI로 추천합니다.
    API 키는 서버 환경변수에서 자동으로 로드됩니다.
    """
    title: str | None = None  # P1 버그 수정: 예외 발생 전 title 미정의 방지
    try:
        data = request.get_json(silent=True) or {}
        url = data.get('url')
        model = data.get('model', DEFAULT_MODEL)

        if not url:
            return jsonify({'error': 'YouTube URL이 필요합니다.'}), 400
        if not content_service.is_youtube_url(url):
            return jsonify({'error': '유효한 YouTube URL을 입력해주세요.'}), 400

        video_id = content_service.get_video_id(url)
        if not video_id:
            return jsonify({'error': '유효하지 않은 YouTube URL입니다.'}), 400

        title = content_service.get_content_title(url) or 'YouTube 영상'

        style_options = current_app.config.get('STYLE_OPTIONS', [])
        style_list = ', '.join([f"{s[0]}({s[1]})" for s in style_options])

        prompt = f"""다음 YouTube 영상 제목을 분석하여 가장 적합한 콘텐츠 스타일을 추천해주세요.

영상 제목: {title}

사용 가능한 스타일: {style_list}

다음 JSON 형식으로만 응답해주세요 (다른 텍스트 없이):
{{
    "style": "추천 스타일 ID",
    "reason": "추천 이유 (20자 이내)",
    "modifiers": {{
        "length": "short|medium|long",
        "tone": "professional|friendly|humorous",
        "emoji": "use|none"
    }}
}}"""

        response = ai_service.create_content(
            prompt,
            model,
            style_prompt=""
        )

        import re
        content = response.get('content', '')
        json_match = re.search(r'\{[\s\S]*\}', content)
        if json_match:
            recommendation = json.loads(json_match.group())
            recommendation['title'] = title
            return jsonify(recommendation)

        return jsonify({
            'style': 'detailed',
            'reason': '기본 추천',
            'modifiers': {'length': 'medium', 'tone': 'professional', 'emoji': 'none'},
            'title': title
        })

    except json.JSONDecodeError:
        return jsonify({
            'style': 'detailed',
            'reason': 'AI 응답 파싱 실패',
            'modifiers': {'length': 'medium', 'tone': 'professional', 'emoji': 'none'},
            'title': title or 'YouTube 영상'
        })
    except Exception as e:
        current_app.logger.error(f"Recommend style failed: {e}")
        return _handle_error_response(str(e))


@blog_bp.route('/api/generate-style', methods=['POST'])
def generate_style():
    """YouTube 제목과 자막을 분석하여 맞춤형 프롬프트를 AI로 생성합니다.
    API 키는 서버 환경변수에서 자동으로 로드됩니다.
    """
    title: str | None = None  # P1 버그 수정: 예외 발생 전 title 미정의 방지
    try:
        data = request.get_json(silent=True) or {}
        url = data.get('url')
        model = data.get('model', DEFAULT_MODEL)

        if not url:
            return jsonify({'error': 'YouTube URL이 필요합니다.'}), 400
        if not content_service.is_youtube_url(url):
            return jsonify({'error': '유효한 YouTube URL을 입력해주세요.'}), 400

        video_id = content_service.get_video_id(url)
        if not video_id:
            return jsonify({'error': '유효하지 않은 YouTube URL입니다.'}), 400

        title = content_service.get_content_title(url) or 'YouTube 영상'

        transcript = content_service.get_transcript(video_id)
        if isinstance(transcript, dict) and transcript.get('error'):
            transcript_preview = "(자막 없음)"
        else:
            # P2 버그 #4 수정: transcript가 dict이면 text 필드 추출
            text = transcript.get('text', '') if isinstance(transcript, dict) else str(transcript or '')
            transcript_preview = text[:500] if text else "(자막 없음)"

        prompt = f"""다음 YouTube 영상의 제목과 자막 일부를 분석하여, 이 영상 콘텐츠에 최적화된 블로그 작성 프롬프트를 생성해주세요.

영상 제목: {title}
자막 미리보기: {transcript_preview}

다음 JSON 형식으로만 응답해주세요:
{{
    "styleName": "이 스타일의 이름 (10자 이내, 예: '기술 분석', '리뷰 정리')",
    "stylePrompt": "AI가 블로그를 작성할 때 사용할 상세 프롬프트 (200-400자)",
    "description": "이 스타일이 왜 이 영상에 적합한지 (30자 이내)"
}}"""

        response = ai_service.create_content(
            prompt,
            model,
            style_prompt=""
        )

        import re
        content = response.get('content', '')
        json_match = re.search(r'\{[\s\S]*\}', content)
        if json_match:
            result = json.loads(json_match.group())
            result['title'] = title
            return jsonify(result)

        return jsonify({
            'error': 'AI 응답을 파싱할 수 없습니다.',
            'title': title
        }), 500

    except json.JSONDecodeError:
        return jsonify({
            'error': 'AI 응답 JSON 파싱 실패',
            'title': title or 'YouTube 영상'
        }), 500
    except Exception as e:
        current_app.logger.error(f"Generate style failed: {e}")
        return _handle_error_response(str(e))


def _generate_main_content(app, content, model, style_prompt, modifiers, style_id=None):
    """스레드에서 메인 콘텐츠를 생성합니다.

    Returns:
        tuple: (result_dict, used_prompt)
    """
    with app.app_context():
        return ai_service.create_content(
            content, model, style_prompt,
            return_prompt=True, modifiers=modifiers,
            style_id=style_id
        )


def _generate_comment_summary(app, comments, model):
    """스레드에서 댓글 요약을 생성합니다. 실패 시 None 반환.

    Args:
        app: Flask 앱 객체
        comments: 댓글 리스트 (최대 50개 사용)
        model: AI 모델 ID

    Returns:
        dict 또는 None: {'content': '...', 'usage': {...}} 또는 실패 시 None
    """
    from prompts.styles.comment_summary import COMMENT_SUMMARY_PROMPT

    with app.app_context():
        try:
            comments_text = '\n'.join(comments[:50])
            comment_content = f"[시청자 댓글]\n{comments_text}"

            result = ai_service.create_content(
                comment_content, model, COMMENT_SUMMARY_PROMPT,
                style_id='comment_summary'
            )
            return result
        except Exception as e:
            current_app.logger.warning(f"댓글 요약 생성 실패 (무시): {e}")
            return None


def _combine_results(main_result, main_prompt, comment_result):
    """메인 결과와 댓글 요약을 결합합니다.

    Args:
        main_result: 메인 AI 생성 결과 dict
        main_prompt: 메인 생성에 사용된 프롬프트
        comment_result: 댓글 요약 결과 dict 또는 None

    Returns:
        tuple: (combined_result, used_prompt)
    """
    if not comment_result:
        return main_result, main_prompt

    import markdown as md_lib

    # 본문 끝에 댓글 요약 추가
    comment_body = comment_result.get('content', '')
    combined_content = main_result.get('content', '') + '\n\n' + comment_body

    # HTML 재생성
    try:
        combined_html = md_lib.markdown(
            combined_content,
            extensions=['tables', 'fenced_code', 'nl2br']
        )
    except Exception:
        combined_html = main_result.get('html', '') + comment_result.get('html', '')

    # 토큰 사용량 합산
    main_usage = main_result.get('usage', {})
    comment_usage = comment_result.get('usage', {})
    combined_usage = {
        'prompt_tokens': main_usage.get('prompt_tokens', 0) + comment_usage.get('prompt_tokens', 0),
        'completion_tokens': main_usage.get('completion_tokens', 0) + comment_usage.get('completion_tokens', 0),
        'total_tokens': main_usage.get('total_tokens', 0) + comment_usage.get('total_tokens', 0),
    }

    combined_result = {
        'title': main_result.get('title', ''),
        'content': combined_content,
        'html': combined_html,
        'usage': combined_usage,
    }

    return combined_result, main_prompt


def _handle_short_content_bypass(transcript_text, style, youtube_title,
                                  raw_transcript, transcript_source, start_time):
    """짧은 콘텐츠 바이패스 체크. 해당 시 Response 반환, 아니면 None."""
    from config import SHORT_CONTENT_THRESHOLD, SHORT_CONTENT_BYPASS_STYLES
    if len(transcript_text) >= SHORT_CONTENT_THRESHOLD or style not in SHORT_CONTENT_BYPASS_STYLES:
        return None

    import markdown as md_lib
    g.skip_usage_decrement = True
    bypass_html = md_lib.markdown(transcript_text, extensions=['tables', 'fenced_code', 'nl2br'])
    elapsed_time = round(time.time() - start_time, 2)
    return jsonify({
        'title': youtube_title,
        'content': transcript_text,
        'html': bypass_html,
        'id': str(uuid.uuid4()),
        'prompt': '',
        'elapsed_time': elapsed_time,
        'youtube_title': youtube_title,
        'transcript': raw_transcript,
        'transcript_source': transcript_source,
        'comment_summary_included': False,
        'bypassed': True,
        'bypass_reason': 'short_content',
        'usage': get_usage_for_response()
    })


def _handle_cache_hit(cache_key, force, youtube_title,
                       raw_transcript, transcript_source, start_time):
    """캐시 히트 체크. 히트 시 Response 반환, 아니면 None."""
    if force:
        return None

    cached = current_app.ai_cache.get(cache_key)
    if not cached:
        return None

    g.skip_usage_decrement = True
    elapsed_time = round(time.time() - start_time, 2)
    return jsonify({
        **cached,
        'id': str(uuid.uuid4()),
        'elapsed_time': elapsed_time,
        'youtube_title': youtube_title,
        'transcript': raw_transcript,
        'transcript_source': transcript_source,
        'cached': True,
        'duplicate_message': '동일 설정으로 이전에 생성된 콘텐츠입니다.',
        'usage': get_usage_for_response()
    })


def _call_ai_with_comments(truncated_content, model, style_prompt, params,
                            comments, transcript_text, max_tokens):
    """AI 호출 + 댓글 병렬 처리. (result, used_prompt, comment_result) 반환."""
    is_auto = model == 'auto'

    if is_auto:
        from config import FALLBACK_CHAIN
        comment_result = None
        if comments:
            combined_content = _build_combined_content(transcript_text, comments)
            truncated_content = content_service.truncate_text(
                f"[영상 자막]\n{combined_content}", max_tokens
            )
        result, used_prompt = ai_service.create_content_with_fallback(
            truncated_content, FALLBACK_CHAIN, style_prompt,
            return_prompt=True, modifiers=params['modifiers'],
            style_id=params['style']
        )
        return result, used_prompt, comment_result

    is_glm = model.startswith('zhipuai/')

    if comments:
        app = current_app._get_current_object()

        if is_glm:
            # GLM 모델: 순차 실행 (글로벌 락 충돌 방지)
            result, used_prompt = ai_service.create_content(
                truncated_content, model, style_prompt,
                return_prompt=True, modifiers=params['modifiers'],
                style_id=params['style']
            )
            comment_result = _generate_comment_summary(app, comments, model)
        else:
            # 병렬 실행
            with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
                main_future = executor.submit(
                    _generate_main_content, app, truncated_content,
                    model, style_prompt, params['modifiers'],
                    style_id=params['style']
                )
                comment_future = executor.submit(
                    _generate_comment_summary, app, comments, model
                )

                result, used_prompt = main_future.result()
                comment_result = comment_future.result()

        result, used_prompt = _combine_results(result, used_prompt, comment_result)
    else:
        # 댓글 없음: 단일 AI 호출
        comment_result = None
        result, used_prompt = ai_service.create_content(
            truncated_content, model, style_prompt,
            return_prompt=True, modifiers=params['modifiers'],
            style_id=params['style']
        )

    return result, used_prompt, comment_result


def _save_and_respond(result, used_prompt, comment_result, cache_key,
                       video_id, params, url, youtube_title,
                       raw_transcript, transcript_source, comments, start_time):
    """캐시 저장 + 히스토리 저장 + JSON 응답 반환."""
    model = params['model']
    modifiers = params['modifiers'] or {}

    current_app.ai_cache.put(
        cache_key, video_id, params['style'], model,
        modifiers.get('length', 'medium'),
        modifiers.get('writing_style', 'conversational'),
        {
            'title': result.get('title', ''),
            'content': result.get('content', ''),
            'html': result.get('html', ''),
            'comment_summary_included': bool(comments and comment_result),
            'prompt': used_prompt,
        }
    )

    elapsed_time = round(time.time() - start_time, 2)

    report_id = str(uuid.uuid4())
    if g.user_id:
        save_history(g.user_id, {
            'id': report_id,
            'url': url,
            'title': result.get('title', youtube_title),
            'style': params['style'],
            'content': result.get('content', ''),
            'html': result.get('html', ''),
            'transcript': raw_transcript,
            'transcript_source': transcript_source,
            'usage': result.get('usage'),
            'elapsed_time': elapsed_time
        })

    # SEO 메타데이터 추출 (blog_seo 스타일만)
    seo = None
    if params['style'] == 'blog_seo':
        seo = ai_service.extract_seo_metadata(result.get('content', ''))

    return jsonify({
        **result,
        "id": report_id,
        "prompt": used_prompt,
        "elapsed_time": elapsed_time,
        "youtube_title": youtube_title,
        "transcript": raw_transcript,
        "transcript_source": transcript_source,
        "comment_summary_included": bool(comments and comment_result),
        "seo": seo,
        "usage": get_usage_for_response()
    })


@blog_bp.route('/generate', methods=['POST'])
@require_auth
@require_usage
def generate():
    """단일 YouTube URL에서 콘텐츠를 생성합니다.
    API 키는 서버 환경변수에서 자동으로 로드됩니다.
    로그인 필수, 하루 5회 제한 적용 (관리자는 무제한).
    """
    try:
        start_time = time.time()
        params = _get_request_data(request)
        url = params['url']

        if not url:
            return jsonify({'error': 'YouTube URL이 필요합니다.'}), 400
        if not content_service.is_youtube_url(url):
            return jsonify({'error': '유효한 YouTube URL을 입력해주세요.'}), 400

        video_id = content_service.get_video_id(url)
        if not video_id:
            return jsonify({'error': '유효하지 않은 YouTube URL입니다.'}), 400

        youtube_title = content_service.get_content_title(url) or 'YouTube 영상'

        transcript_text, comments, error, raw_transcript, transcript_source = _fetch_youtube_content(video_id)
        if error:
            return jsonify({'error': error}), 400

        max_tokens = get_model_max_tokens(params['model'])
        main_content = f"[영상 자막]\n{transcript_text}"
        truncated_content = content_service.truncate_text(main_content, max_tokens)

        # 짧은 콘텐츠 바이패스
        bypass_resp = _handle_short_content_bypass(
            transcript_text, params['style'], youtube_title,
            raw_transcript, transcript_source, start_time
        )
        if bypass_resp:
            return bypass_resp

        # 캐시 체크
        from services.cache_service import AICacheService
        force = (request.get_json(silent=True) or {}).get('force', False)
        modifiers = params['modifiers'] or {}
        cache_key = AICacheService.make_key(
            video_id, params['style'], params['model'],
            modifiers.get('length', 'medium'),
            modifiers.get('writing_style', 'conversational')
        )
        cache_resp = _handle_cache_hit(
            cache_key, force, youtube_title,
            raw_transcript, transcript_source, start_time
        )
        if cache_resp:
            return cache_resp

        # AI 호출 (auto/GLM/병렬 분기)
        style_prompt = _get_style_prompt(params['style'], params['custom_prompt'])
        result, used_prompt, comment_result = _call_ai_with_comments(
            truncated_content, params['model'], style_prompt, params,
            comments, transcript_text, max_tokens
        )

        # 캐시 저장 + 히스토리 + 응답
        return _save_and_respond(
            result, used_prompt, comment_result, cache_key,
            video_id, params, url, youtube_title,
            raw_transcript, transcript_source, comments, start_time
        )

    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        current_app.logger.error(f"Generate failed: {e}")
        return _handle_error_response(str(e))


@blog_bp.route('/regenerate', methods=['POST'])
@require_auth
@require_usage
def regenerate():
    """기존 콘텐츠를 새로운 스타일로 재생성합니다.
    API 키는 서버 환경변수에서 자동으로 로드됩니다.
    로그인 필수, 하루 5회 제한 적용 (관리자는 무제한).
    """
    try:
        params = _get_request_data(request)
        content = params['content']

        if not content:
            return jsonify({'error': '재생성할 콘텐츠가 없습니다'}), 400

        style_prompt = _get_style_prompt(params['style'])
        result, used_prompt = ai_service.create_content(
            content,
            params['model'],
            style_prompt,
            return_prompt=True
        )

        return jsonify({**result, "prompt": used_prompt, "usage": get_usage_for_response()})

    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        current_app.logger.error(f"Regenerate failed: {e}")
        return _handle_error_response(str(e))


def _process_single_url(app, url, model, style, modifiers, custom_prompt):
    """배치 처리에서 단일 URL을 처리하는 헬퍼 함수입니다.
    API 키는 서버 환경변수에서 자동으로 로드됩니다.
    """
    with app.app_context():
        try:
            current_app.logger.info(f"Processing URL: {url}")

            if not content_service.is_youtube_url(url):
                return {
                    'success': False,
                    'url': url,
                    'title': 'URL 오류',
                    'error': '유효한 YouTube URL이 아닙니다.'
                }

            video_id = content_service.get_video_id(url)
            if not video_id:
                return {
                    'success': False,
                    'url': url,
                    'title': 'YouTube 영상',
                    'error': '유효하지 않은 YouTube URL입니다'
                }

            title = content_service.get_content_title(url) or 'YouTube 영상'
            current_app.logger.info(f"Content title: {title}")

            transcript_text, comments, error, raw_transcript, transcript_source = _fetch_youtube_content(video_id)
            if error:
                return {
                    'success': False,
                    'url': url,
                    'title': title,
                    'error': error
                }

            # 배치 처리: 기존 방식(자막+댓글 합성)으로 처리
            content = _build_combined_content(transcript_text, comments)
            max_tokens = get_model_max_tokens(model)
            content = content_service.truncate_text(content, max_tokens)
            style_prompt = _get_style_prompt(style, custom_prompt)

            result, used_prompt = ai_service.create_content(
                content, model, style_prompt,
                return_prompt=True, modifiers=modifiers,
                style_id=style
            )

            return {
                'success': True,
                'url': url,
                'title': result.get('title', title),
                'content': result.get('content', ''),
                'html': result.get('html', ''),
                'prompt': used_prompt,
                'transcript_source': transcript_source
            }

        except Exception as e:
            current_app.logger.error(f"Error processing URL {url}: {e}")
            return {
                'success': False,
                'url': url,
                'title': '오류 발생',
                'error': _sanitize_error_for_client(f'처리 중 오류 발생: {str(e)}')
            }


@blog_bp.route('/generate-batch', methods=['POST'])
@require_auth
def generate_batch():
    """여러 URL을 배치로 처리합니다.
    API 키는 서버 환경변수에서 자동으로 로드됩니다.
    로그인 필수, 하루 5회 제한 적용 (배치 전체가 1회로 계산, 관리자는 무제한).
    """
    from services.usage.usage_service import UsageService

    try:
        # 원자적 사용량 체크 + 차감 (Race Condition 방지)
        can_use, usage = UsageService.try_consume_atomic(g.user_id)
        if not can_use:
            return jsonify({
                'error': '오늘 사용 가능 횟수를 모두 소진했습니다. 내일 다시 시도해주세요.',
                'code': 'USAGE_LIMIT_EXCEEDED',
                'usage': usage
            }), 429

        current_app.logger.info("Batch generate request received")

        data = request.get_json()
        current_app.logger.info(f"Request data: {data}")

        if not data:
            current_app.logger.error("No JSON data received")
            return jsonify({'error': 'JSON 데이터가 제공되지 않았습니다'}), 400

        urls = data.get('urls', [])
        model = data.get('model', DEFAULT_MODEL)
        style = data.get('style', DEFAULT_STYLE)
        modifiers = data.get('modifiers')
        custom_prompt = data.get('customPrompt')

        current_app.logger.info(f"URLs to process: {urls}, Model: {model}, Style: {style}")

        if not urls or not isinstance(urls, list):
            return jsonify({'error': 'URL 목록이 제공되지 않았습니다'}), 400
        if len(urls) > MAX_BATCH_URLS:
            return jsonify({'error': f'최대 {MAX_BATCH_URLS}개의 URL만 처리할 수 있습니다'}), 400

        app = current_app._get_current_object()
        results = [None] * len(urls)
        combined_content = []

        # GLM-4.7은 동시성 제한으로 순차 처리 필요
        is_sequential_model = model == 'zhipuai/GLM-4.7'

        if is_sequential_model:
            current_app.logger.info(f"Starting to process {len(urls)} URLs sequentially (GLM-4.7)")
            for i, url in enumerate(urls):
                try:
                    result = _process_single_url(app, url, model, style, modifiers, custom_prompt)
                    results[i] = result
                    current_app.logger.info(f"Completed processing URL {i + 1}: {result.get('success', False)}")

                    if result['success'] and isinstance(result.get('content', ''), str):
                        combined_content.append(result['content'])
                except Exception as e:
                    current_app.logger.error(f"Exception for URL {i + 1}: {e}")
                    results[i] = {
                        'success': False,
                        'url': url,
                        'title': '오류 발생',
                        'error': _sanitize_error_for_client(str(e))
                    }
        else:
            current_app.logger.info(f"Starting to process {len(urls)} URLs concurrently")

            with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_BATCH_WORKERS) as executor:
                future_to_index = {
                    executor.submit(
                        _process_single_url, app, url, model, style,
                        modifiers, custom_prompt
                    ): i for i, url in enumerate(urls)
                }

                try:
                    for future in concurrent.futures.as_completed(future_to_index, timeout=300):
                        index = future_to_index[future]
                        try:
                            result = future.result(timeout=120)
                            results[index] = result
                            current_app.logger.info(f"Completed processing URL {index + 1}: {result.get('success', False)}")

                            if result['success'] and isinstance(result.get('content', ''), str):
                                combined_content.append(result['content'])
                        except concurrent.futures.TimeoutError:
                            current_app.logger.error(f"Timeout for URL {index + 1}")
                            results[index] = {
                                'success': False,
                                'url': urls[index],
                                'title': '시간 초과',
                                'error': '[타임아웃] 처리 시간이 초과되었습니다.'
                            }
                        except Exception as e:
                            current_app.logger.error(f"Exception in future for URL {index + 1}: {e}")
                            results[index] = {
                                'success': False,
                                'url': urls[index],
                                'title': '오류 발생',
                                'error': _sanitize_error_for_client(str(e))
                            }
                except concurrent.futures.TimeoutError:
                    current_app.logger.error("Batch overall timeout (300s)")
                    for future, index in future_to_index.items():
                        if results[index] is None:
                            future.cancel()
                            results[index] = {
                                'success': False,
                                'url': urls[index],
                                'title': '시간 초과',
                                'error': '[타임아웃] 배치 전체 처리 시간이 초과되었습니다.'
                            }

        url_to_result = {result['url']: result for result in results if result}
        ordered_results = [
            url_to_result.get(url, {'success': False, 'url': url, 'error': '처리 실패'})
            for url in urls
        ]

        final_combined_content = "\n\n=== 다음 콘텐츠 ===\n\n".join(combined_content) if combined_content else ""
        success_count = sum(1 for r in ordered_results if r.get('success'))
        fail_count = len(ordered_results) - success_count

        current_app.logger.info(f"Batch processing completed. Success: {success_count}, Failed: {fail_count}")

        # 사용량은 try_consume_atomic()으로 이미 차감됨
        updated_usage = usage

        # P2 버그 #7 수정: 배치 히스토리 저장 (N+1 → 배치 INSERT)
        # 배치에서는 transcript, usage, elapsed_time이 None (P3 #13 문서화)
        if g.user_id:
            histories_to_save = []
            for result in ordered_results:
                if result.get('success'):
                    report_id = str(uuid.uuid4())
                    result['id'] = report_id
                    histories_to_save.append({
                        'id': report_id,
                        'url': result.get('url'),
                        'title': result.get('title'),
                        'style': style,
                        'content': result.get('content', ''),
                        'html': result.get('html', ''),
                        'transcript': None,
                        'usage': None,
                        'elapsed_time': None
                    })

            # 배치 INSERT (1회 DB 호출)
            if histories_to_save:
                from services.supabase_service import get_supabase
                supabase = get_supabase()
                if supabase:
                    try:
                        def sanitize_text(text, max_length=50000):
                            """입력 텍스트 검증 및 길이 제한"""
                            if not isinstance(text, str):
                                return ""
                            return text[:max_length]

                        batch_data = [{
                            'user_id': g.user_id,
                            'report_id': sanitize_text(h.get('id', ''), 100),
                            'url': sanitize_text(h.get('url', ''), 500),
                            'title': sanitize_text(h.get('title', ''), 500),
                            'style': sanitize_text(h.get('style', ''), 50),
                            'content': sanitize_text(h.get('content', ''), 100000),
                            'html': sanitize_text(h.get('html', ''), 200000),
                            'transcript': sanitize_text(h.get('transcript', ''), 10000),
                            'usage': h.get('usage'),
                            'elapsed_time': h.get('elapsed_time')
                        } for h in histories_to_save]
                        supabase.table('ie_histories').insert(batch_data).execute()
                    except Exception as e:
                        current_app.logger.warning(f"배치 히스토리 저장 실패: {e}")

        return jsonify({
            'success': True,
            'results': ordered_results,
            'content': final_combined_content,
            'total_processed': len(urls),
            'successful': success_count,
            'failed': fail_count,
            'usage': updated_usage
        })

    except ValueError as e:
        current_app.logger.error(f"ValueError in batch generate: {e}")
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        current_app.logger.error(f"Batch generate failed: {e}", exc_info=True)
        return _handle_error_response(f'배치 처리 중 오류: {str(e)}')


MAX_MERGED_URLS = 5


@blog_bp.route('/api/generate-merged', methods=['POST'])
@require_auth
@require_usage
def generate_merged():
    """여러 YouTube URL의 자막을 합쳐 하나의 통합 콘텐츠를 생성합니다.
    2~5개 URL 지원, AI 호출 1회, 사용량 1회 차감.
    """
    try:
        start_time = time.time()
        params = _get_request_data(request)
        urls = params['urls']

        if len(urls) < 2:
            return jsonify({'error': '합쳐서 생성은 최소 2개 URL이 필요합니다.'}), 400
        if len(urls) > MAX_MERGED_URLS:
            return jsonify({'error': f'최대 {MAX_MERGED_URLS}개 URL까지 합칠 수 있습니다.'}), 400

        # URL 유효성 검사
        for url in urls:
            if not content_service.is_youtube_url(url):
                return jsonify({'error': f'유효하지 않은 YouTube URL: {url}'}), 400

        # 병렬로 자막+댓글 추출
        app = current_app._get_current_object()
        video_data = []  # (url, video_id, title, transcript, comments, source)

        def _fetch_one(url):
            with app.app_context():
                vid = content_service.get_video_id(url)
                if not vid:
                    return {'url': url, 'error': '유효하지 않은 YouTube URL'}
                title = content_service.get_content_title(url) or 'YouTube 영상'
                transcript_text, comments, error, _, source = _fetch_youtube_content(vid)
                if error:
                    return {'url': url, 'error': error, 'title': title}
                return {
                    'url': url, 'video_id': vid, 'title': title,
                    'transcript': transcript_text, 'comments': comments,
                    'transcript_source': source,
                }

        with concurrent.futures.ThreadPoolExecutor(max_workers=min(MAX_MERGED_URLS, len(urls))) as executor:
            futures = {executor.submit(_fetch_one, u): u for u in urls}
            for future in concurrent.futures.as_completed(futures):
                video_data.append(future.result())

        # URL 순서 복원
        url_order = {u: i for i, u in enumerate(urls)}
        video_data.sort(key=lambda d: url_order.get(d['url'], 99))

        # 실패 URL 확인
        successes = [d for d in video_data if 'transcript' in d]
        if len(successes) < 2:
            errors = [f"{d['title']}: {d['error']}" for d in video_data if 'error' in d]
            return jsonify({
                'error': f'자막 추출 성공이 2개 미만입니다. 실패: {"; ".join(errors)}'
            }), 400

        # 토큰 예산 분배 후 합성
        max_tokens = get_model_max_tokens(params['model'])
        prompt_overhead = 4000
        available_tokens = max_tokens - prompt_overhead
        per_url_tokens = available_tokens // len(successes)

        merged_parts = []
        all_comments = []
        source_videos = []

        for d in successes:
            truncated = content_service.truncate_text(d['transcript'], per_url_tokens)
            merged_parts.append(f"[영상 {len(merged_parts) + 1}: {d['title']}]\n{truncated}")
            if d['comments']:
                all_comments.extend(d['comments'][:10])
            source_videos.append({
                'url': d['url'],
                'title': d['title'],
                'transcript_source': d['transcript_source'],
            })

        merged_content = '\n\n'.join(merged_parts)
        if all_comments:
            comments_text = '\n'.join(all_comments[:30])
            merged_content += f"\n\n[시청자 댓글 (종합)]\n{comments_text}"

        # AI 호출 1회
        style_prompt = _get_style_prompt(params['style'], params['custom_prompt'])
        result, used_prompt = ai_service.create_content(
            merged_content, params['model'], style_prompt,
            return_prompt=True, modifiers=params['modifiers'],
            style_id=params['style']
        )

        elapsed_time = round(time.time() - start_time, 2)
        report_id = str(uuid.uuid4())

        # 히스토리 저장 (첫 번째 URL 기준)
        if g.user_id:
            save_history(g.user_id, {
                'id': report_id,
                'url': urls[0],
                'title': result.get('title', '통합 분석'),
                'style': params['style'],
                'content': result.get('content', ''),
                'html': result.get('html', ''),
                'transcript': None,
                'usage': result.get('usage'),
                'elapsed_time': elapsed_time,
            })

        # SEO 메타데이터
        seo = None
        if params['style'] == 'blog_seo':
            seo = ai_service.extract_seo_metadata(result.get('content', ''))

        return jsonify({
            **result,
            'id': report_id,
            'prompt': used_prompt,
            'elapsed_time': elapsed_time,
            'source_videos': source_videos,
            'merged': True,
            'seo': seo,
            'usage': get_usage_for_response(),
        })

    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        current_app.logger.error(f"Generate merged failed: {e}")
        return _handle_error_response(str(e))


@blog_bp.route('/api/mindmap', methods=['POST'])
@require_auth
@require_usage
def generate_mindmap():
    """기존 콘텐츠를 마인드맵 형식의 마크다운으로 변환합니다.
    API 키는 서버 환경변수에서 자동으로 로드됩니다.
    로그인 필수, 하루 5회 제한 적용 (관리자는 무제한).
    """
    try:
        start_time = time.time()
        data = request.get_json(silent=True) or {}
        content = data.get('content')
        model = data.get('model', DEFAULT_MODEL)

        if not content:
            return jsonify({'error': '마인드맵으로 변환할 콘텐츠가 필요합니다.'}), 400

        # MINDMAP_PROMPT 가져오기
        style_prompts = current_app.config.get('STYLE_PROMPTS', {})
        mindmap_prompt = style_prompts.get('mindmap', '')

        if not mindmap_prompt:
            return jsonify({'error': '마인드맵 프롬프트가 설정되지 않았습니다.'}), 500

        # 콘텐츠 길이 제한 (토큰 절약)
        max_tokens = get_model_max_tokens(model)
        truncated_content = content_service.truncate_text(content, min(max_tokens, 50000))

        result = ai_service.create_content(
            truncated_content,
            model,
            mindmap_prompt
        )

        elapsed_time = round(time.time() - start_time, 2)

        # 마인드맵용 마크다운 콘텐츠 반환
        return jsonify({
            'success': True,
            'markdown': result.get('content', ''),
            'elapsed_time': elapsed_time,
            'usage': get_usage_for_response()
        })

    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        current_app.logger.error(f"Mindmap generation failed: {e}")
        return _handle_error_response(str(e))


@blog_bp.route('/api/generate-multi', methods=['POST'])
@require_auth
@require_usage
def generate_multi():
    """하나의 URL을 여러 스타일로 동시에 생성합니다.
    사용량: 멀티 생성 전체 = 1회 차감.
    """
    try:
        start_time = time.time()
        data = request.get_json(silent=True) or {}
        url = data.get('url')
        model = data.get('model', DEFAULT_MODEL)
        styles = data.get('styles', [])

        if not url:
            return jsonify({'error': 'YouTube URL이 필요합니다.'}), 400
        if not content_service.is_youtube_url(url):
            return jsonify({'error': '유효한 YouTube URL을 입력해주세요.'}), 400
        if not styles or not isinstance(styles, list) or len(styles) < 1:
            return jsonify({'error': '최소 1개 이상의 스타일을 선택해주세요.'}), 400
        if len(styles) > 5:
            return jsonify({'error': '최대 5개 스타일까지 선택할 수 있습니다.'}), 400

        video_id = content_service.get_video_id(url)
        if not video_id:
            return jsonify({'error': '유효하지 않은 YouTube URL입니다.'}), 400

        youtube_title = content_service.get_content_title(url) or 'YouTube 영상'

        transcript_text, comments, error, raw_transcript, transcript_source = _fetch_youtube_content(video_id)
        if error:
            return jsonify({'error': error}), 400

        max_tokens = get_model_max_tokens(model)
        main_content = _build_combined_content(transcript_text, comments) if comments else f"[영상 자막]\n{transcript_text}"
        truncated_content = content_service.truncate_text(main_content, max_tokens)

        # 유효 스타일 필터링
        style_prompts_dict = current_app.config.get('STYLE_PROMPTS', {})
        valid_styles = [s for s in styles if s in style_prompts_dict]
        if not valid_styles:
            return jsonify({'error': '유효한 스타일이 없습니다.'}), 400

        # 병렬 생성
        app = current_app._get_current_object()
        results = []

        def _gen_for_style(style_id):
            with app.app_context():
                try:
                    sp = style_prompts_dict.get(style_id, '')
                    result = ai_service.create_content(
                        truncated_content, model, sp,
                        style_id=style_id
                    )
                    result['style'] = style_id
                    return result
                except Exception as e:
                    return {'style': style_id, 'error': str(e)}

        is_glm = model.startswith('zhipuai/')
        if is_glm:
            # GLM: 순차 실행
            for style_id in valid_styles:
                results.append(_gen_for_style(style_id))
        else:
            with concurrent.futures.ThreadPoolExecutor(max_workers=min(3, len(valid_styles))) as executor:
                futures = {executor.submit(_gen_for_style, s): s for s in valid_styles}
                for future in concurrent.futures.as_completed(futures):
                    results.append(future.result())

        # 스타일 순서 정렬
        style_order = {s: i for i, s in enumerate(valid_styles)}
        results.sort(key=lambda r: style_order.get(r.get('style', ''), 99))

        elapsed_time = round(time.time() - start_time, 2)

        return jsonify({
            'success': True,
            'results': results,
            'youtube_title': youtube_title,
            'transcript_source': transcript_source,
            'elapsed_time': elapsed_time,
            'usage': get_usage_for_response()
        })

    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        current_app.logger.error(f"Generate multi failed: {e}")
        return _handle_error_response(str(e))


@blog_bp.route('/api/generate-fusion', methods=['POST'])
@require_auth
@check_usage
def generate_fusion():
    """퓨전 생성: N개 URL → 융합 1편"""
    data = request.get_json()
    urls = data.get('urls', [])
    style_id = data.get('style', 'blog_seo')
    model = data.get('model', '')
    modifiers = data.get('modifiers', {})
    enable_web_research = data.get('enable_web_research', True)
    enable_deep_comments = data.get('enable_deep_comments', True)

    if not urls or len(urls) < 2:
        return jsonify({'error': '[입력 오류] 퓨전 분석은 최소 2개 URL이 필요합니다'}), 400
    if len(urls) > 5:
        return jsonify({'error': '[입력 오류] 퓨전 분석은 최대 5개 URL까지 가능합니다'}), 400
    if not model:
        return jsonify({'error': '[입력 오류] 모델을 선택해주세요'}), 400

    try:
        result = fusion_service.generate_fusion(
            urls=urls,
            style_id=style_id,
            model=model,
            modifiers=modifiers,
            enable_web_research=enable_web_research,
            enable_deep_comments=enable_deep_comments
        )

        # 사용량 차감 (1회)
        if hasattr(g, 'usage') and g.usage:
            from services.usage import UsageService
            UsageService.decrement(g.usage.get('user_id'))

        return jsonify(result)

    except ValueError as e:
        return jsonify({'error': f'[입력 오류] {str(e)}'}), 400
    except Exception as e:
        current_app.logger.error('퓨전 생성 실패: %s', e, exc_info=True)
        return _handle_error_response(e)


@blog_bp.route('/api/playlist-videos', methods=['POST'])
@require_auth
def playlist_videos():
    """채널 또는 재생목록 URL에서 영상 목록을 추출합니다."""
    try:
        data = request.get_json(silent=True) or {}
        url = data.get('url', '')
        max_results = min(int(data.get('maxResults', 10)), 50)

        if not url:
            return jsonify({'error': 'URL이 필요합니다.'}), 400

        if content_service.is_playlist_url(url):
            result = content_service.get_playlist_videos(url, max_results)
        elif content_service.is_channel_url(url):
            result = content_service.get_channel_videos(url, max_results)
        else:
            return jsonify({'error': '유효한 채널 또는 재생목록 URL이 아닙니다.'}), 400

        if 'error' in result:
            return jsonify(result), 400

        return jsonify(result)

    except Exception as e:
        current_app.logger.error(f"Playlist videos failed: {e}")
        return jsonify({'error': '영상 목록을 가져올 수 없습니다.'}), 500


@blog_bp.route('/api/export/docx', methods=['POST'])
@require_auth
def export_docx():
    """마크다운 콘텐츠를 DOCX 파일로 변환하여 반환합니다."""
    try:
        data = request.get_json(silent=True) or {}
        title = data.get('title', 'AI 생성 결과')
        content = data.get('content', '')

        if not content:
            return jsonify({'error': '변환할 콘텐츠가 없습니다.'}), 400

        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 기본 스타일 설정
        style = doc.styles['Normal']
        style.font.name = 'Malgun Gothic'
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.5

        # 제목
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 마크다운 → docx 변환
        _markdown_to_docx(doc, content)

        # BytesIO에 저장
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        safe_title = re_module.sub(r'[^\w\s가-힣-]', '', title)[:30].strip() or 'content'
        filename = f'{safe_title}.docx'

        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except ImportError:
        return jsonify({'error': 'python-docx 패키지가 설치되지 않았습니다.'}), 500
    except Exception as e:
        current_app.logger.error(f"DOCX export failed: {e}")
        return jsonify({'error': 'DOCX 변환 중 오류가 발생했습니다.'}), 500


def _markdown_to_docx(doc, markdown_text):
    """마크다운 텍스트를 docx Document에 변환하여 추가합니다."""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    lines = markdown_text.split('\n')
    i = 0
    in_list = False
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 빈 줄
        if not stripped:
            if in_table and table_rows:
                _add_table_to_docx(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        # 테이블 행 (| ... | 형식)
        if stripped.startswith('|') and stripped.endswith('|'):
            # 구분선 행 (|---|---|) 건너뛰기
            if re_module.match(r'^\|[\s\-:]+\|$', stripped):
                i += 1
                continue
            in_table = True
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
            continue

        # 테이블 종료
        if in_table and table_rows:
            _add_table_to_docx(doc, table_rows)
            table_rows = []
            in_table = False

        # 헤딩
        if stripped.startswith('###'):
            text = stripped.lstrip('#').strip()
            doc.add_heading(_strip_markdown_formatting(text), level=3)
        elif stripped.startswith('##'):
            text = stripped.lstrip('#').strip()
            doc.add_heading(_strip_markdown_formatting(text), level=2)
        elif stripped.startswith('#'):
            text = stripped.lstrip('#').strip()
            doc.add_heading(_strip_markdown_formatting(text), level=2)

        # 수평선
        elif stripped in ('---', '***', '___'):
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(6)

        # 인용문
        elif stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            p = doc.add_paragraph(style='Quote') if 'Quote' in [s.name for s in doc.styles] else doc.add_paragraph()
            _add_formatted_text(p, text)

        # 비순서 목록
        elif stripped.startswith(('- ', '* ', '+ ')):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)

        # 순서 목록
        elif re_module.match(r'^\d+\.\s', stripped):
            text = re_module.sub(r'^\d+\.\s', '', stripped).strip()
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, text)

        # 일반 텍스트
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)

        i += 1

    # 남은 테이블 처리
    if in_table and table_rows:
        _add_table_to_docx(doc, table_rows)


def _strip_markdown_formatting(text):
    """마크다운 인라인 서식을 제거합니다."""
    text = re_module.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re_module.sub(r'\*(.+?)\*', r'\1', text)
    text = re_module.sub(r'`(.+?)`', r'\1', text)
    return text


def _add_formatted_text(paragraph, text):
    """마크다운 인라인 서식을 docx Run으로 변환하여 추가합니다."""
    # 볼드, 이탤릭, 코드 패턴 분리
    pattern = re_module.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last_end = 0

    for match in pattern.finditer(text):
        # 매치 전 일반 텍스트
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        if match.group(2):  # 볼드
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # 이탤릭
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # 코드
            run = paragraph.add_run(match.group(4))
            run.font.name = 'Consolas'

        last_end = match.end()

    # 나머지 텍스트
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _add_table_to_docx(doc, rows):
    """테이블 데이터를 docx Table로 추가합니다."""
    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = _strip_markdown_formatting(cell_text)
                # 첫 행(헤더) 볼드
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


@blog_bp.route('/api/cache/stats', methods=['GET'])
@require_auth
def api_cache_stats():
    """AI 결과 캐시 통계를 반환합니다."""
    stats = current_app.ai_cache.get_stats()
    return jsonify(stats)


@blog_bp.route('/api/cache/ai', methods=['DELETE'])
@require_auth
def api_clear_ai_cache():
    """AI 결과 캐시를 삭제합니다. videoId가 있으면 해당 영상만."""
    data = request.get_json(silent=True) or {}
    video_id = data.get('videoId')
    deleted = current_app.ai_cache.clear(video_id)
    return jsonify({'success': True, 'deleted': deleted})


@blog_bp.route('/generate-stream', methods=['POST'])
@require_auth
def generate_stream():
    """SSE 스트리밍으로 콘텐츠를 생성합니다.
    @require_usage 데코레이터 사용 불가 (generator 응답) → 수동 사용량 관리.
    GLM/auto 모델은 스트리밍 미지원 → 비스트리밍 /generate 사용 권장.
    """
    from services.usage.usage_service import UsageService
    from services.supabase_service import is_supabase_enabled
    import markdown as md_lib

    try:
        params = _get_request_data(request)
        url = params['url']

        if not url:
            return jsonify({'error': 'YouTube URL이 필요합니다.'}), 400
        if not content_service.is_youtube_url(url):
            return jsonify({'error': '유효한 YouTube URL을 입력해주세요.'}), 400

        video_id = content_service.get_video_id(url)
        if not video_id:
            return jsonify({'error': '유효하지 않은 YouTube URL입니다.'}), 400

        # 사용량 체크 (수동)
        user_id = getattr(g, 'user_id', None)
        if is_supabase_enabled() and user_id:
            can_use, usage = UsageService.check_can_use(user_id)
            if not can_use:
                return jsonify({
                    'error': '오늘 사용 가능 횟수를 모두 소진했습니다.',
                    'code': 'USAGE_LIMIT_EXCEEDED',
                    'usage': usage
                }), 429

        youtube_title = content_service.get_content_title(url) or 'YouTube 영상'
        transcript_text, comments, error, raw_transcript, transcript_source = _fetch_youtube_content(video_id)
        if error:
            return jsonify({'error': error}), 400

        max_tokens = get_model_max_tokens(params['model'])
        main_content = f"[영상 자막]\n{transcript_text}"
        if comments:
            main_content = _build_combined_content(transcript_text, comments)
        truncated_content = content_service.truncate_text(main_content, max_tokens)

        style_prompt = _get_style_prompt(params['style'], params['custom_prompt'])
        model = params['model']

        app = current_app._get_current_object()

        def generate_sse():
            with app.app_context():
                try:
                    # meta 이벤트
                    meta = json.dumps({
                        'type': 'meta',
                        'youtube_title': youtube_title,
                        'transcript_source': transcript_source
                    }, ensure_ascii=False)
                    yield f"data: {meta}\n\n"

                    # 토큰 스트리밍
                    full_content = ''
                    for token in ai_service.create_content_stream(
                        truncated_content, model, style_prompt,
                        modifiers=params['modifiers'], style_id=params['style']
                    ):
                        if token is None:
                            break
                        full_content += token
                        token_event = json.dumps({
                            'type': 'token',
                            'content': token
                        }, ensure_ascii=False)
                        yield f"data: {token_event}\n\n"

                    # 완료: 제목/본문 분리 + HTML 변환
                    title = youtube_title
                    body = full_content
                    lines = full_content.split('\n')
                    if lines and lines[0].startswith('#'):
                        title = lines[0].lstrip('#').strip()
                        body = '\n'.join(lines[1:]).strip()

                    try:
                        html = md_lib.markdown(body, extensions=['tables', 'fenced_code', 'nl2br'])
                    except Exception:
                        html = f"<pre>{body}</pre>"

                    # 사용량 차감 (성공 시)
                    if is_supabase_enabled() and user_id:
                        UsageService.decrement(user_id)

                    done_event = json.dumps({
                        'type': 'done',
                        'title': title,
                        'content': body,
                        'html': html,
                        'youtube_title': youtube_title,
                        'transcript_source': transcript_source,
                    }, ensure_ascii=False)
                    yield f"data: {done_event}\n\n"

                except Exception as e:
                    error_event = json.dumps({
                        'type': 'error',
                        'message': str(e)
                    }, ensure_ascii=False)
                    yield f"data: {error_event}\n\n"

        return Response(
            stream_with_context(generate_sse()),
            mimetype='text/event-stream',
            headers={
                'Cache-Control': 'no-cache',
                'X-Accel-Buffering': 'no',
            }
        )

    except Exception as e:
        current_app.logger.error(f"Generate stream setup failed: {e}")
        return _handle_error_response(str(e))
