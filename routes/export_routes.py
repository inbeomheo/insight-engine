"""
내보내기 라우트 — DOCX 변환
"""
import io
import re as re_module

from flask import request, jsonify, current_app, send_file

from routes.blog_routes import blog_bp
from services.data.supabase_service import require_auth
from utils.responses import handle_error


@blog_bp.route('/api/export/docx', methods=['POST'])
@require_auth
def export_docx():
    """마크다운 콘텐츠를 DOCX 파일로 변환하여 반환합니다."""
    try:
        data = request.get_json(silent=True) or {}
        title = data.get('title', 'AI 생성 결과')
        content = data.get('content', '')

        if not content:
            return jsonify({'error': '변환할 콘텐츠가 없습니다.'}), 400

        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        include_toc = data.get('include_toc', False)

        doc = Document()

        # 기본 스타일 설정
        style = doc.styles['Normal']
        style.font.name = 'Malgun Gothic'
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.5

        # 제목
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 목차 삽입 (include_toc=true 시)
        if include_toc:
            _insert_toc(doc, content)

        # 마크다운 → docx 변환
        _markdown_to_docx(doc, content)

        # BytesIO에 저장
        buffer = io.BytesIO()
        doc.save(buffer)

        # Content-Length 계산 (다운로드 진행률 표시용)
        content_length = buffer.seek(0, 2)
        buffer.seek(0)

        safe_title = re_module.sub(r'[^\w\s가-힣-]', '', title)[:30].strip() or 'content'
        filename = f'{safe_title}.docx'

        response = send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        response.headers['Content-Length'] = content_length
        return response

    except ImportError:
        return jsonify({'error': 'python-docx 패키지가 설치되지 않았습니다.'}), 500
    except Exception as e:
        current_app.logger.error(f"DOCX export failed: {e}")
        return jsonify({'error': 'DOCX 변환 중 오류가 발생했습니다.'}), 500


def _extract_headings(markdown_text):
    """마크다운에서 헤딩을 추출합니다. [(level, text), ...]"""
    headings = []
    for line in markdown_text.split('\n'):
        stripped = line.strip()
        if stripped.startswith('#'):
            # 헤딩 레벨 계산
            level = 0
            for ch in stripped:
                if ch == '#':
                    level += 1
                else:
                    break
            if 1 <= level <= 4 and len(stripped) > level:
                text = stripped[level:].strip()
                text = re_module.sub(r'\*\*(.+?)\*\*', r'\1', text)
                text = re_module.sub(r'\*(.+?)\*', r'\1', text)
                text = re_module.sub(r'`(.+?)`', r'\1', text)
                if text:
                    headings.append((level, text))
    return headings


def _insert_toc(doc, markdown_text):
    """마크다운 헤딩 기반 목차를 DOCX에 삽입합니다."""
    from docx.shared import Pt

    headings = _extract_headings(markdown_text)
    if not headings:
        return

    # 목차 제목
    toc_heading = doc.add_heading('목차', level=2)

    # 최소 레벨 기준으로 들여쓰기
    min_level = min(h[0] for h in headings)

    for level, text in headings:
        indent = '  ' * (level - min_level)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f'{indent}{"─ " if level > min_level else "■ "}{text}')
        run.font.size = Pt(10)

    # 구분선
    doc.add_paragraph()


def _markdown_to_docx(doc, markdown_text):
    """마크다운 텍스트를 docx Document에 변환하여 추가합니다."""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    lines = markdown_text.split('\n')
    i = 0
    in_list = False
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 빈 줄
        if not stripped:
            if in_table and table_rows:
                _add_table_to_docx(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        # 테이블 행 (| ... | 형식)
        if stripped.startswith('|') and stripped.endswith('|'):
            # 구분선 행 (|---|---|) 건너뛰기
            if re_module.match(r'^\|[\s\-:]+\|$', stripped):
                i += 1
                continue
            in_table = True
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
            continue

        # 테이블 종료
        if in_table and table_rows:
            _add_table_to_docx(doc, table_rows)
            table_rows = []
            in_table = False

        # 헤딩
        if stripped.startswith('###'):
            text = stripped.lstrip('#').strip()
            doc.add_heading(_strip_markdown_formatting(text), level=3)
        elif stripped.startswith('##'):
            text = stripped.lstrip('#').strip()
            doc.add_heading(_strip_markdown_formatting(text), level=2)
        elif stripped.startswith('#'):
            text = stripped.lstrip('#').strip()
            doc.add_heading(_strip_markdown_formatting(text), level=2)

        # 수평선
        elif stripped in ('---', '***', '___'):
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(6)

        # 인용문
        elif stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            p = doc.add_paragraph(style='Quote') if 'Quote' in [s.name for s in doc.styles] else doc.add_paragraph()
            _add_formatted_text(p, text)

        # 비순서 목록
        elif stripped.startswith(('- ', '* ', '+ ')):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)

        # 순서 목록
        elif re_module.match(r'^\d+\.\s', stripped):
            text = re_module.sub(r'^\d+\.\s', '', stripped).strip()
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, text)

        # 일반 텍스트
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)

        i += 1

    # 남은 테이블 처리
    if in_table and table_rows:
        _add_table_to_docx(doc, table_rows)


def _strip_markdown_formatting(text):
    """마크다운 인라인 서식을 제거합니다."""
    text = re_module.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re_module.sub(r'\*(.+?)\*', r'\1', text)
    text = re_module.sub(r'`(.+?)`', r'\1', text)
    return text


def _add_formatted_text(paragraph, text):
    """마크다운 인라인 서식을 docx Run으로 변환하여 추가합니다."""
    # 볼드, 이탤릭, 코드 패턴 분리
    pattern = re_module.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last_end = 0

    for match in pattern.finditer(text):
        # 매치 전 일반 텍스트
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        if match.group(2):  # 볼드
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # 이탤릭
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # 코드
            run = paragraph.add_run(match.group(4))
            run.font.name = 'Consolas'

        last_end = match.end()

    # 나머지 텍스트
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


@blog_bp.route('/api/export/epub', methods=['POST'])
@require_auth
def export_epub():
    """마크다운 콘텐츠를 EPUB 전자책 파일로 변환하여 반환합니다."""
    try:
        data = request.get_json(silent=True) or {}
        title = data.get('title', 'AI 생성 결과')
        content = data.get('content', '')
        author = data.get('author', '')

        if not content:
            return jsonify({'error': '변환할 콘텐츠가 없습니다.'}), 400

        from services.export.epub_service import create_epub
        epub_bytes = create_epub(title, content, author)

        buffer = io.BytesIO(epub_bytes)
        buffer.seek(0)

        safe_title = re_module.sub(r'[^\w\s가-힣-]', '', title)[:30].strip() or 'content'
        filename = f'{safe_title}.epub'

        return send_file(
            buffer,
            mimetype='application/epub+zip',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        current_app.logger.error(f"EPUB export failed: {e}")
        return jsonify({'error': 'EPUB 변환 중 오류가 발생했습니다.'}), 500


@blog_bp.route('/api/export/markdown', methods=['POST'])
@require_auth
def export_markdown():
    """마크다운 파일로 내보냅니다."""
    try:
        data = request.get_json(silent=True) or {}
        title = data.get('title', 'AI 생성 결과')
        content = data.get('content', '')
        if not content:
            return jsonify({'error': '변환할 콘텐츠가 없습니다.'}), 400

        frontmatter = data.get('frontmatter')

        from services.export.export_service import export_markdown as _export_md
        buffer = _export_md(title, content, frontmatter=frontmatter)
        safe_title = re_module.sub(r'[^\w\s가-힣-]', '', title)[:30].strip() or 'content'

        return send_file(buffer, mimetype='text/markdown', as_attachment=True, download_name=f'{safe_title}.md')
    except Exception as e:
        current_app.logger.error(f"MD export failed: {e}")
        return jsonify({'error': '마크다운 내보내기 실패'}), 500


@blog_bp.route('/api/export/txt', methods=['POST'])
@require_auth
def export_txt():
    """텍스트 파일로 내보냅니다."""
    try:
        data = request.get_json(silent=True) or {}
        title = data.get('title', 'AI 생성 결과')
        content = data.get('content', '')
        if not content:
            return jsonify({'error': '변환할 콘텐츠가 없습니다.'}), 400

        line_ending = data.get('line_ending', 'unix')
        if line_ending not in ('unix', 'windows'):
            line_ending = 'unix'

        from services.export.export_service import export_txt as _export_txt
        buffer = _export_txt(title, content, line_ending=line_ending)
        safe_title = re_module.sub(r'[^\w\s가-힣-]', '', title)[:30].strip() or 'content'

        return send_file(buffer, mimetype='text/plain', as_attachment=True, download_name=f'{safe_title}.txt')
    except Exception as e:
        current_app.logger.error(f"TXT export failed: {e}")
        return jsonify({'error': '텍스트 내보내기 실패'}), 500


@blog_bp.route('/api/export/zip', methods=['POST'])
@require_auth
def export_zip():
    """ZIP 패키지 (DOCX+MD+TXT+meta.json)로 내보냅니다."""
    try:
        data = request.get_json(silent=True) or {}
        title = data.get('title', 'AI 생성 결과')
        content = data.get('content', '')
        if not content:
            return jsonify({'error': '변환할 콘텐츠가 없습니다.'}), 400

        from services.export.export_service import export_zip as _export_zip
        buffer = _export_zip(title, content)
        safe_title = re_module.sub(r'[^\w\s가-힣-]', '', title)[:30].strip() or 'content'

        return send_file(buffer, mimetype='application/zip', as_attachment=True, download_name=f'{safe_title}.zip')
    except Exception as e:
        current_app.logger.error(f"ZIP export failed: {e}")
        return jsonify({'error': 'ZIP 내보내기 실패'}), 500


@blog_bp.route('/api/export/slides', methods=['POST'])
@require_auth
def export_slides():
    """마크다운 콘텐츠를 Reveal.js 슬라이드 HTML로 변환합니다."""
    try:
        data = request.get_json(silent=True) or {}
        content = data.get('content', '')
        theme = data.get('theme', 'black')
        title = data.get('title', '슬라이드')

        if not content:
            return jsonify({'error': '변환할 콘텐츠가 없습니다.'}), 400

        from services.media.slide_service import convert_to_slides
        html = convert_to_slides(content, theme)

        buf = io.BytesIO(html.encode('utf-8'))
        buf.seek(0)
        safe_title = re_module.sub(r'[^\w\s가-힣-]', '', title)[:30].strip() or 'slides'

        return send_file(buf, mimetype='text/html', as_attachment=True, download_name=f'{safe_title}.html')

    except ValueError as e:
        return handle_error(str(e))
    except Exception as e:
        current_app.logger.error(f"Slide export failed: {e}")
        return jsonify({'error': '슬라이드 변환 중 오류가 발생했습니다.'}), 500


@blog_bp.route('/api/export/srt', methods=['POST'])
@require_auth
def export_srt():
    """자막 세그먼트를 SRT 포맷으로 변환하여 다운로드합니다."""
    try:
        data = request.get_json(silent=True) or {}
        segments = data.get('segments', [])
        title = data.get('title', '자막')

        if not segments:
            return jsonify({'error': '변환할 자막 세그먼트가 없습니다.'}), 400

        srt_lines = []
        for idx, seg in enumerate(segments, 1):
            start = _seconds_to_srt_time(seg.get('start', 0))
            end = _seconds_to_srt_time(seg.get('end', seg.get('start', 0) + seg.get('duration', 0)))
            text = seg.get('text', '').strip()
            if text:
                srt_lines.append(f'{idx}\n{start} --> {end}\n{text}\n')

        srt_content = '\n'.join(srt_lines)
        buf = io.BytesIO(srt_content.encode('utf-8'))
        buf.seek(0)
        safe_title = re_module.sub(r'[^\w\s가-힣-]', '', title)[:30].strip() or 'subtitle'

        return send_file(buf, mimetype='text/srt', as_attachment=True, download_name=f'{safe_title}.srt')

    except Exception as e:
        current_app.logger.error(f"SRT export failed: {e}")
        return jsonify({'error': 'SRT 변환 중 오류가 발생했습니다.'}), 500


def _seconds_to_srt_time(seconds: float) -> str:
    """초를 SRT 타임코드 형식(HH:MM:SS,mmm)으로 변환합니다."""
    seconds = max(0, float(seconds))
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f'{h:02d}:{m:02d}:{s:02d},{ms:03d}'


@blog_bp.route('/api/export/infographic', methods=['POST'])
@require_auth
def export_infographic():
    """콘텐츠를 HTML 인포그래픽으로 변환합니다."""
    try:
        data = request.get_json(silent=True) or {}
        title = data.get('title', '인포그래픽')
        content = data.get('content', '')
        if not content:
            return jsonify({'error': '변환할 콘텐츠가 없습니다.'}), 400

        from services.media.infographic_service import generate_infographic
        html_content = generate_infographic(content, title)

        buffer = io.BytesIO(html_content.encode('utf-8'))
        buffer.seek(0)
        safe_title = re_module.sub(r'[^\w\s가-힣-]', '', title)[:30].strip() or 'infographic'

        return send_file(buffer, mimetype='text/html', as_attachment=True,
                         download_name=f'{safe_title}_infographic.html')
    except Exception as e:
        current_app.logger.error(f"Infographic export failed: {e}")
        return jsonify({'error': '인포그래픽 생성 실패'}), 500


@blog_bp.route('/api/export/card-news', methods=['POST'])
@require_auth
def export_card_news():
    """콘텐츠를 카드뉴스 SVG 세트 ZIP으로 내보냅니다."""
    try:
        import zipfile
        data = request.get_json(silent=True) or {}
        title = data.get('title', '카드뉴스')
        content = data.get('content', '')
        if not content:
            return jsonify({'error': '변환할 콘텐츠가 없습니다.'}), 400

        from services.media.card_news_service import generate_card_news
        cards = generate_card_news(content, title)
        if not cards:
            return jsonify({'error': '카드뉴스로 변환할 포인트가 없습니다.'}), 400

        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            for i, svg in enumerate(cards):
                zf.writestr(f'card_{i+1:02d}.svg', svg)
        buffer.seek(0)

        safe_title = re_module.sub(r'[^\w\s가-힣-]', '', title)[:30].strip() or 'card_news'

        return send_file(buffer, mimetype='application/zip', as_attachment=True,
                         download_name=f'{safe_title}_cards.zip')
    except Exception as e:
        current_app.logger.error(f"Card news export failed: {e}")
        return jsonify({'error': '카드뉴스 생성 실패'}), 500


@blog_bp.route('/api/export/summary-card', methods=['POST'])
@require_auth
def export_summary_card():
    """제목+핵심 포인트를 요약 카드 SVG로 내보냅니다."""
    try:
        data = request.get_json(silent=True) or {}
        title = data.get('title', '요약')
        key_points = data.get('key_points', [])
        if not key_points:
            return jsonify({'error': '핵심 포인트가 필요합니다.'}), 400

        from services.content.summary_card_service import generate_summary_card
        svg = generate_summary_card(title, key_points)

        buffer = io.BytesIO(svg.encode('utf-8'))
        buffer.seek(0)
        safe_title = re_module.sub(r'[^\w\s가-힣-]', '', title)[:30].strip() or 'summary'

        return send_file(buffer, mimetype='image/svg+xml', as_attachment=True,
                         download_name=f'{safe_title}_card.svg')
    except Exception as e:
        current_app.logger.error(f"Summary card export failed: {e}")
        return jsonify({'error': '요약 카드 생성 실패'}), 500


@blog_bp.route('/api/export/code-image', methods=['POST'])
@require_auth
def export_code_image():
    """코드 스니펫을 carbon.sh 스타일 HTML로 내보냅니다."""
    try:
        data = request.get_json(silent=True) or {}
        code = data.get('code', '')
        language = data.get('language', 'python')
        title = data.get('title', '')
        if not code:
            return jsonify({'error': '코드가 필요합니다.'}), 400

        from services.media.code_image_service import generate_code_image
        html_content = generate_code_image(code, language, title)

        buffer = io.BytesIO(html_content.encode('utf-8'))
        buffer.seek(0)
        safe_title = re_module.sub(r'[^\w\s가-힣-]', '', title)[:30].strip() or 'code'

        return send_file(buffer, mimetype='text/html', as_attachment=True,
                         download_name=f'{safe_title}_code.html')
    except Exception as e:
        current_app.logger.error(f"Code image export failed: {e}")
        return jsonify({'error': '코드 이미지 생성 실패'}), 500


@blog_bp.route('/api/export/newsletter-html', methods=['POST'])
@require_auth
def export_newsletter_html():
    """콘텐츠를 이메일 뉴스레터 HTML로 변환합니다."""
    try:
        data = request.get_json(silent=True) or {}
        title = data.get('title', 'Newsletter')
        content = data.get('content', '')
        if not content:
            return jsonify({'error': '변환할 콘텐츠가 없습니다.'}), 400

        from services.content.newsletter_template_service import convert_to_newsletter
        html = convert_to_newsletter(content, title)

        return jsonify({'success': True, 'html': html})
    except ValueError as e:
        return handle_error(str(e))
    except Exception as e:
        current_app.logger.error(f"Newsletter export failed: {e}")
        return jsonify({'error': '뉴스레터 변환 실패'}), 500


@blog_bp.route('/api/export/interactive-report', methods=['POST'])
@require_auth
def export_interactive_report():
    """콘텐츠를 인터랙티브 HTML 보고서로 변환합니다."""
    try:
        data = request.get_json(silent=True) or {}
        title = data.get('title', '보고서')
        content = data.get('content', '')
        if not content:
            return jsonify({'error': '변환할 콘텐츠가 없습니다.'}), 400

        from services.content.interactive_report_service import generate_interactive_report
        html = generate_interactive_report(content, title)

        return jsonify({'success': True, 'html': html})
    except ValueError as e:
        return handle_error(str(e))
    except Exception as e:
        current_app.logger.error(f"Interactive report export failed: {e}")
        return jsonify({'error': '인터랙티브 보고서 변환 실패'}), 500


def _add_table_to_docx(doc, rows):
    """테이블 데이터를 docx Table로 추가합니다."""
    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = _strip_markdown_formatting(cell_text)
                # 첫 행(헤더) 볼드
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


@blog_bp.route('/api/export/html', methods=['POST'])
@require_auth
def export_html():
    """마크다운 콘텐츠를 독립 HTML 파일로 변환합니다.

    인라인 CSS 포함, 외부 의존성 없는 self-contained HTML.
    Body: {"title": "...", "content": "...", "html": "..."}
    """
    try:
        data = request.get_json(silent=True) or {}
        title = data.get('title', 'AI 생성 결과')
        html_content = data.get('html', '') or data.get('content', '')
        dark_mode = data.get('dark_mode', False)

        if not html_content:
            return jsonify({'error': '변환할 콘텐츠가 없습니다.'}), 400

        dark_css = ""
        if dark_mode:
            dark_css = """
@media (prefers-color-scheme: dark) {
body { background: #1a1a2e; color: #e0e0e0; }
h1 { border-bottom-color: #444; }
h2 { color: #d1d5db; }
h3 { color: #9ca3af; }
blockquote { background: #1e293b; color: #93c5fd; border-left-color: #60a5fa; }
code { background: #374151; color: #f9fafb; }
th, td { border-color: #4b5563; }
th { background: #1f2937; }
a { color: #60a5fa; }
.meta { color: #9ca3af; }
}"""

        standalone_html = f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<style>
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; max-width: 800px; margin: 0 auto; padding: 2rem; line-height: 1.8; color: #1a1a1a; }}
h1 {{ font-size: 1.75rem; border-bottom: 2px solid #e5e7eb; padding-bottom: 0.5rem; }}
h2 {{ font-size: 1.4rem; margin-top: 2rem; color: #374151; }}
h3 {{ font-size: 1.15rem; color: #4b5563; }}
p {{ margin: 0.75rem 0; }}
ul, ol {{ padding-left: 1.5rem; }}
li {{ margin: 0.25rem 0; }}
blockquote {{ border-left: 4px solid #3b82f6; margin: 1rem 0; padding: 0.5rem 1rem; background: #eff6ff; color: #1e40af; }}
code {{ background: #f3f4f6; padding: 0.15rem 0.4rem; border-radius: 4px; font-size: 0.9em; }}
pre {{ background: #1f2937; color: #e5e7eb; padding: 1rem; border-radius: 8px; overflow-x: auto; }}
pre code {{ background: transparent; padding: 0; }}
table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; }}
th, td {{ border: 1px solid #d1d5db; padding: 0.5rem 0.75rem; text-align: left; }}
th {{ background: #f9fafb; font-weight: 600; }}
a {{ color: #2563eb; }}
.meta {{ color: #6b7280; font-size: 0.85rem; margin-bottom: 1.5rem; }}
{dark_css}
</style>
</head>
<body>
<h1>{title}</h1>
<p class="meta">Insight Engine으로 생성됨</p>
{html_content}
</body>
</html>"""

        buf = io.BytesIO(standalone_html.encode('utf-8'))
        safe_title = re_module.sub(r'[^\w\s가-힣-]', '', title)[:50].strip() or 'export'
        return send_file(
            buf,
            mimetype='text/html',
            as_attachment=True,
            download_name=f'{safe_title}.html',
        )

    except Exception as e:
        return handle_error(e, 'HTML 내보내기')
